import os
import tkinter as tk
from tkinter import filedialog
import stat
import shutil
import sys
from appJar import gui
# create a GUI variable called app
developer1 = "Leido por AiLector - Ronel Cruz [ronel.cruz.a8@gmail.com]"
developer = "Si quieres remover este texto de todas tus tus lectura, contactame - Ronel Cruz [ronel.cruz.a8@gmail.com]"
formTitle = "AiLector"
formBGcolor = "orange"
formFontSize = 18
formMainLabel = "Bienvenido a AiLector v2.3"
defVal = "N/a"
formLoadedFileFormat = "[Formato] ["+defVal+"]"
formLoadedFileSize = "[Tamaño] ["+defVal+"]"
formLoadedFilePath = "[Ruta] ["+defVal+"]"
formStatus = "Seleccione un Archivo"
app = gui(formTitle, "600x200")
def press(button):
    setStatus("Por favor espere")
    if button == "Cancel":
        app.stop()
    if button == "Seleccionar Imagen o PDF":
       file = selectFile()
       formStatus = "Archivo Ca       rgado"
       setStatus(formStatus)
       extension = os.path.splitext(file)[1]
       extensionOk = False
       allowedImageFormats = [".png", ".jpg", ".jpeg"]
       
       allowedDocsFormats = [".pdf"]
       if extension in allowedImageFormats:
           extensionOk = True
           setStatus("L E Y E N D O    I M A G E N    P N G ")
           print("------------------>   L E Y E N D O    I M A G E N    P N G ")
           setStatus("Leyendo " + file)
           print("Leyendo " + file)
           import pytesseract
           pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
           from PIL import Image
           from pytesseract import image_to_string
           texto = image_to_string(Image.open(file))
           print("## LECTURA COMPLETADA")
           setStatus(" LECTURA COMPLETADA")
       if extension in allowedDocsFormats:
           print("------------------>   L E Y E N D O     D O C U M E N T O     P D F  ")
           setStatus("L E Y E N D O     D O C U M E N T O     P D F  ")
           extensionOk = True
           # Import libraries
           from PIL import Image   
           import pytesseract 
           pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
           import sys 
           from pdf2image import convert_from_path
           #convert_from_path(poppler_path=r'C:\Program
           #Files\plopper\poppler-0.68.0\bin')
           clear_dir('tmp')
           os.mkdir('tmp')

            # Path of the pdf
           PDF_file = file
  
           ''' 
           Part #1 : Converting PDF to images 
           '''

            # Store all the pages of the PDF in a variable
            # pages = convert_from_path(PDF_file, 500)
            # pages = convert_from_path(PDF_file, poppler_path='C:\Program
            # Files\plopper\poppler-0.68.0\bin')
           pages = convert_from_path(PDF_file, poppler_path= r'C:\Program Files\plopper\poppler-0.68.0\bin')
            # Counter to store images of each page of PDF to image
           image_counter = 1
  
            # Iterate through all the pages stored above
           for page in pages: 
  
                # Declaring filename for each page of PDF as JPG
                # For each page, filename will be:
                # PDF page 1 -> page_1.jpg
                # PDF page 2 -> page_2.jpg
                # PDF page 3 -> page_3.jpg
                # ....
                # PDF page n -> page_n.jpg
                filename = "/tmp/page_" + str(image_counter) + ".jpg"
      
                # Save the image of the page in system
                page.save(filename, 'JPEG') 
  
                # Increment the counter to update filename
                image_counter = image_counter + 1
  
           ''' 
           Part #2 - Recognizing text from the images using OCR 
           '''
            # Variable to get count of total number of pages
           filelimit = image_counter - 1
  
            # Creating a text file to write the output
           outfile = "/tmp/out_text.txt"
  
            # Open the file in append mode so that
            # All contents of all images are added to the same file
           f = open(outfile, "a") 
  
            # Iterate from 1 to total number of pages
           for i in range(1, filelimit + 1): 
  
                # Set filename to recognize text from
                # Again, these files will be:
                # page_1.jpg
                # page_2.jpg
                # ....
                # page_n.jpg
                filename = "/tmp/page_" + str(i) + ".jpg"
          
                # Recognize the text as string in image using pytesserct
                text = str(((pytesseract.image_to_string(Image.open(filename))))) 
  
                # The recognized text is stored in variable text
                # Any string processing may be applied on text
                # Here, basic formatting has been done:
                # In many PDFs, at line ending, if a word can't
                # be written fully, a 'hyphen' is added.
                # The rest of the word is written in the next line
                # Eg: This is a sample text this word here GeeksF-
                # orGeeks is half on first line, remaining on next.
                # To remove this, we replace every '-\n' to ''.
                texto = text.replace('-\n', '')     
  
                # Finally, write the processed text to the file.
                f.write(text) 
  
            # Close the file after writing all the text.
           f.close() 
           from pathlib import Path
           texto = Path('/tmp/out_text.txt').read_text()
           clear_dir('tmp')
           print("## LECTURA COMPLETADA")
           setStatus(" LECTURA COMPLETADA")

       def valid_xml_char_ordinal(c):
            codepoint = ord(c)
            # conditions ordered by presumed frequency
            return (
                0x20 <= codepoint <= 0xD7FF or
                codepoint in (0x9, 0xA, 0xD) or
                0xE000 <= codepoint <= 0xFFFD or
                0x10000 <= codepoint <= 0x10FFFF
                )

       if extensionOk == True:
            #setStatus("G U A R D A N D O     A R C H I V O     T X T  ")
            print("------------------>   G U A R D A N D O     A R C H I V O     T X T  ")
            outputExt = app.getOptionBox("Formato de salida")
            exxt = "File"
            exxt2 = "*"+outputExt
            
            file = filedialog.asksaveasfile(mode='w', defaultextension=outputExt, filetypes=((exxt, exxt2),("All Files", "*.*")))
            if file is None: # asksaveasfile return `None` if dialog closed with "cancel".
                return
            if outputExt == ".txt":
                file.write(texto)
                file.write(developer1)
                file.write(developer)
                file.close() # `()` was missing.

            if outputExt == ".docx":
                from docx import Document
                document = Document()
                cleaned_string1 = ''.join(c for c in texto if valid_xml_char_ordinal(c))
                cleaned_string2 = ''.join(c for c in developer1 if valid_xml_char_ordinal(c))
                cleaned_string3 = ''.join(c for c in developer if valid_xml_char_ordinal(c))
                document.add_paragraph(cleaned_string1)
                document.add_paragraph(cleaned_string2)
                document.add_paragraph(cleaned_string3)
                document.save(file.name)

            
            messageBox("Archivo completado")
            setStatus("Archivo completado")
            setStatus("A B R I E N D O    A R C H I V O     T X T  ")
            print("------------------>   A B R I E N D O    A R C H I V O     T X T  ")
            openFile(file)
       if extension == False:
           #dosomething
           setStatus("Formato Incorrecto")
           messageBox("El formato del archivo que has seleccionado es "+extension+" y este por el momento no es admitido.")

                 








def c():
    print("______________________________________________________________________________________________")


def menu0():
    print("         S E L E C C I O N E      U N A      O P C I O N ")
    c()
    menu = "1-Nueva lectura ","2-Salir "
    return menu
def selectFile():
    root = tk.Tk()
    root.withdraw()
    return filedialog.askopenfilename()
def read():
    #file = selectFile()
    import os.path
    extension = os.path.splitext(file)[1]
    extensionOk = False
    allowedImageFormats = [".png", ".jpg", ".jpeg"]
    allowedDocsFormats = [".pdf"]
         
    if extension in allowedImageFormats:
        extensionOk = True
        print("------------------>   L E Y E N D O    I M A G E N    P N G ")
        print("Leyendo " + file)
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        from PIL import Image
        from pytesseract import image_to_string
        texto = image_to_string(Image.open(file))
        print("## LECTURA COMPLETADA")
    if extension in allowedDocsFormats:
        print("------------------>   L E Y E N D O     D O C U M E N T O     P D F  ")
        extensionOk = True
        # Import libraries
        from PIL import Image 
        import pytesseract 
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        import sys 
        from pdf2image import convert_from_path
        
        
        #convert_from_path(poppler_path=r'C:\Program
        #Files\plopper\poppler-0.68.0\bin')
        clear_dir('tmp')
        os.mkdir('tmp')

        # Path of the pdf
        PDF_file = file
  
        ''' 
        Part #1 : Converting PDF to images 
        '''

        # Store all the pages of the PDF in a variable
        # pages = convert_from_path(PDF_file, 500)
        # pages = convert_from_path(PDF_file, poppler_path='C:\Program
        # Files\plopper\poppler-0.68.0\bin')
        pages = convert_from_path(PDF_file, poppler_path= r'C:\Program Files\plopper\poppler-0.68.0\bin')
        # Counter to store images of each page of PDF to image
        image_counter = 1
  
        # Iterate through all the pages stored above
        for page in pages: 
  
            # Declaring filename for each page of PDF as JPG
            # For each page, filename will be:
            # PDF page 1 -> page_1.jpg
            # PDF page 2 -> page_2.jpg
            # PDF page 3 -> page_3.jpg
            # ....
            # PDF page n -> page_n.jpg
            filename = "/tmp/page_" + str(image_counter) + ".jpg"
      
            # Save the image of the page in system
            page.save(filename, 'JPEG') 
  
            # Increment the counter to update filename
            image_counter = image_counter + 1
  
        ''' 
        Part #2 - Recognizing text from the images using OCR 
        '''
        # Variable to get count of total number of pages
        filelimit = image_counter - 1
  
        # Creating a text file to write the output
        outfile = "/tmp/out_text.txt"
  
        # Open the file in append mode so that
        # All contents of all images are added to the same file
        f = open(outfile, "a") 
  
        # Iterate from 1 to total number of pages
        for i in range(1, filelimit + 1): 
  
            # Set filename to recognize text from
            # Again, these files will be:
            # page_1.jpg
            # page_2.jpg
            # ....
            # page_n.jpg
            filename = "/tmp/page_" + str(i) + ".jpg"
          
            # Recognize the text as string in image using pytesserct
            text = str(((pytesseract.image_to_string(Image.open(filename))))) 
  
            # The recognized text is stored in variable text
            # Any string processing may be applied on text
            # Here, basic formatting has been done:
            # In many PDFs, at line ending, if a word can't
            # be written fully, a 'hyphen' is added.
            # The rest of the word is written in the next line
            # Eg: This is a sample text this word here GeeksF-
            # orGeeks is half on first line, remaining on next.
            # To remove this, we replace every '-\n' to ''.
            texto = text.replace('-\n', '')     
  
            # Finally, write the processed text to the file.
            f.write(text) 
  
        # Close the file after writing all the text.
        f.close() 
        from pathlib import Path
        texto = Path('/tmp/out_text.txt').read_text()
        clear_dir('tmp')
    
    if extensionOk == True:
        print("------------------>   G U A R D A N D O     A R C H I V O     T X T  ")
        file = filedialog.asksaveasfile(mode='w', defaultextension=".txt", filetypes=(("TXT file", "*.txt"),("All Files", "*.*")))
        if file is None: # asksaveasfile return `None` if dialog closed with "cancel".
            return
        file.write(texto)
        file.close() # `()` was missing.
        print("------------------>   A B R I E N D O    A R C H I V O     T X T  ")
        openFile(file)
    else:
        c()
        print("Tipo de archivo no admitido")
        c()

def clear_dir(path_):
    if os.path.exists(path_):
        shutil.rmtree(path_)
def openFile(file):
    import subprocess
    subprocess.call(['cmd.exe', '/c', file.name])
def analize(x):
    if x == '2':
        sys.exit()
    if x == '1':
        print("Nueva Lectura >")
        read()

#salir = False;
#while not salir:
#    print(menu0())
#    comando = input(">")
#    analize(comando)




def updateInfo():
    #if formStatus == "["+defVal+"]":
    x=app
    #x.setLabel("format", formLoadedFileFormat)
    #x.setLabel("size", formLoadedFileSize)
    #x.setLabel("path", formLoadedFilePath)
    
def createForm(x):
    availableOutputFormats = [".docx", ".txt"]
    x.setBg(formBGcolor)
    x.setFont(formFontSize)
    x.addLabel("title", formMainLabel)
    #x.addLabel("format", formLoadedFileFormat)
    #x.addLabel("size", formLoadedFileSize)
    #x.addLabel("path", formLoadedFilePath)
    x.addLabel("status", formStatus)
    x.setLabelBg("title", "red")   
    app.addLabelOptionBox("Formato de salida", availableOutputFormats)
    app.addButtons(["Seleccionar Imagen o PDF", "Cancel"], press)
    #app.openBox(title=None, dirName=None, fileTypes=None, asFile=False, parent=None, multiple=False, mode='r')   
    updateInfo() 
    x.go()

def messageBox(mensaje):
    app.infoBox("Informacion", mensaje, parent=None)

def setStatus(status):
    app.setLabel("status", status)




createForm(app)